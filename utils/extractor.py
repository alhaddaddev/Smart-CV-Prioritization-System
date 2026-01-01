import os
import re
import pdfplumber
from docx import Document
import easyocr
import spacy

nlp = spacy.load("en_core_web_sm")

_easy_reader = None

def get_easyocr_reader():
    global _easy_reader
    if _easy_reader is None:
        _easy_reader = easyocr.Reader(["en"], gpu=False)
    return _easy_reader

# ---------------- TEXT EXTRACTION ----------------

def extract_text_from_pdf(path):
    try:
        text = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
        return "\n".join(text)
    except Exception as e:
        print("PDF ERROR:", e)
        return ""

def extract_text_from_docx(path):
    try:
        doc = Document(path)
        texts = []

        # Paragraphs
        for p in doc.paragraphs:
            if p.text.strip():
                texts.append(p.text)

        # Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        texts.append(cell.text)

        return "\n".join(texts)
    except Exception as e:
        print("DOCX ERROR:", e)
        return ""

def extract_text_from_image(path):
    try:
        reader = get_easyocr_reader()
        result = reader.readtext(
            path,
            detail=0,
            paragraph=True
        )
        return "\n".join(result)
    except Exception as e:
        print("OCR ERROR:", e)
        return ""

# Unified extractor
def extract_text(path):
    ext = os.path.splitext(path)[1].lower()

    if ext == ".pdf":
        return extract_text_from_pdf(path)

    if ext == ".docx":
        return extract_text_from_docx(path)

    if ext in (".png", ".jpg", ".jpeg"):
        return extract_text_from_image(path)

    return ""

# ---------------- CLEANING ----------------

def clean_text(text):
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"[^\x00-\x7F]+", " ", text)
    text = re.sub(r"[,:;|()\[\]{}]", " ", text)
    text = text.strip()

    return text

# ---------------- NLP PHRASES ----------------

def extract_nlp_phrases(text, max_phrases):
    text = clean_text(text)
    if not text:
        return []

    doc = nlp(text)
    phrases = set()

    def valid_phrase(p):
        if len(p) < 2:
            return False
        if any(ch in p for ch in "|\\"):
            return False
        return True

    # Noun chunks (skills, roles, education)
    for chunk in doc.noun_chunks:
        p = chunk.text.lower().strip()
        if valid_phrase(p):
            phrases.add(p)

    # Named entities (tools, orgs, degrees)
    for ent in doc.ents:
        p = ent.text.lower().strip()
        if valid_phrase(p):
            phrases.add(p)

    return list(phrases)[:max_phrases]

# ---------------- UI INSIGHTS ----------------

def extract_ui_insights(text, max_items):
    text = clean_text(text)
    if not text:
        return []

    doc = nlp(text)

    allowed_labels = {
        "ORG",
        "PRODUCT",
        "WORK_OF_ART",
        "GPE",
        "EVENT"
    }

    insights = set()

    for ent in doc.ents:
        if ent.label_ in allowed_labels:
            value = ent.text.strip()
            if len(value) > 2:
                insights.add(value)

    return sorted(list(insights))[:max_items]
